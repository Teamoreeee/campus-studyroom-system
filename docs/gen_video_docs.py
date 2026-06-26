#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""把《演示视频分镜与台词.md》转成 Word，并拆成两份：
   1) 演示视频分镜脚本.docx —— 只含分镜（镜头/时间/画面/备注），不含旁白
   2) 演示视频台词稿.docx   —— 只含台词（旁白），按镜头/时间组织
用法：python gen_video_docs.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

FONT = "微软雅黑"
BLUE = RGBColor(0x1E, 0x3A, 0x8A)      # 深蓝标题
BLUE2 = RGBColor(0x25, 0x63, 0xEB)     # 科技蓝
GRAY = RGBColor(0x1F, 0x29, 0x37)      # 正文深灰

# ── 数据：从 md 整理出的 21 个镜头 ───────────────────────────────
TITLE = "校园自习室预约系统 — 演示视频"
SUBINFO = [
    "视频时长：8–10 分钟",
    "建议分辨率：1920×1080（16:9）",
    "录制方式：屏幕录制 + 配音旁白",
    "配音人：建议何展霖主配，后端同学补配技术细节（可选）",
]

TIME_TABLE = [
    ("片头", "00:00 – 00:30", "封面 + 团队介绍", "30 秒"),
    ("项目背景", "00:30 – 01:30", "痛点 + 目标 + 需求", "60 秒"),
    ("架构设计", "01:30 – 03:30", "技术选型 + 微服务 + 4+1 视图", "120 秒"),
    ("数据库与安全", "03:30 – 04:30", "数据库设计 + 安全架构", "60 秒"),
    ("AI 智能模块", "04:30 – 06:00", "智能推荐 + RAG 客服 + AI 辅助开发", "90 秒"),
    ("部署验证", "06:00 – 07:30", "Docker Compose 部署 + 全链路验证", "90 秒"),
    ("总结", "07:30 – 09:00", "团队分工 + 亮点 + 展望 + 片尾", "90 秒"),
]

# 每个镜头: (序号, 标题, 时间, [画面要点], 旁白, 备注)
SHOTS = [
    (1, "片头", "00:00 – 00:30",
     ["黑场淡入，出现系统 Logo 或首页截图", "叠加文字：「校园自习室预约系统」", "切换到 PPT 封面页"],
     "大家好，我们是 CampusStudio 团队。今天为大家介绍我们的课程设计项目——校园自习室预约系统。本项目基于 Spring Cloud Alibaba 微服务架构，结合 AI 智能推荐与 RAG 智能客服，致力于解决高校自习室资源管理中的痛点问题。",
     "语速适中，营造正式感"),
    (2, "项目背景", "00:30 – 01:00",
     ["PPT 第 3 页「项目背景」", "可配一张校园自习室占座的真实照片"],
     "在高校日常学习场景中，自习室资源管理面临几大痛点：高峰期一座难求，学生常常需要逐层找座；占座不学现象普遍，造成资源浪费；考勤统计依赖人工，效率低且易出错；同时，学生缺乏智能化的座位推荐手段。",
     "照片可网络搜索「大学自习室占座」通用图"),
    (3, "项目目标与需求", "01:00 – 01:30",
     ["PPT 第 4 页「项目目标」", "PPT 第 5 页「需求分析」快速切换"],
     "针对这些痛点，我们设定了五大目标：资源数字化、预约智能化、考勤自动化、管理精细化、架构云原生化。系统覆盖七大业务模块，支持学生、管理员、超级管理员三类角色，并围绕创建预约、签到、AI 推荐等八个核心用例展开设计。",
     "两页切换要快，不要停留太久"),
    (4, "技术选型", "01:30 – 02:00",
     ["PPT 第 6 页「技术选型」", "突出显示技术栈表格"],
     "在技术选型上，后端采用 Java 17 + Spring Boot 3.2.5 + Spring Cloud Alibaba 2023.0.1.2，前端采用 Vue 3 + TypeScript + Element Plus。数据层使用 MySQL 8.0，并兼容国产达梦 8 数据库。缓存用 Redis，服务治理用 Nacos，网关采用 Spring Cloud Gateway，整体支持 Docker 与 Kubernetes 云原生部署。",
     "读技术栈时语速稍快，体现专业性"),
    (5, "架构选型论证", "02:00 – 02:20",
     ["PPT 第 7 页「架构选型论证」"],
     "我们对比了微服务与单体架构，以及 Spring Cloud Alibaba 与 Netflix 方案。考虑到系统涉及七个业务领域，且需要体现微服务架构能力，最终选择微服务 + Spring Cloud Alibaba，以满足功能实现和国产化信创的双重要求。",
     "强调「信创」这个加分点"),
    (6, "微服务拆分", "02:20 – 02:45",
     ["PPT 第 8 页「微服务拆分」", "切换到 IDEA 项目结构或架构图"],
     "系统拆分为七个微服务：认证服务 auth、用户服务 user、自习室服务 room、预约服务 reservation、考勤服务 attendance、AI 服务 ai，以及统一的 API 网关 gateway。每个服务职责清晰，独立开发、独立部署。",
     "可在 IDEA 中展示 backend 目录下各服务模块"),
    (7, "4+1 视图 — 场景视图", "02:45 – 03:05",
     ["PPT 第 9 页「场景视图」", "或展示用例图"],
     "架构设计采用 4+1 视图方法。以学生预约自习室为主线：学生登录后浏览自习室，选择座位和时间段提交预约，系统进行冲突检测并创建记录，最后在预约时段内签到签退。",
     "场景视图是串联其他视图的线索"),
    (8, "4+1 视图 — 逻辑视图与开发视图", "03:05 – 03:30",
     ["PPT 第 10 页「逻辑视图」", "PPT 第 11 页「开发视图 + 进程视图」", "可展示项目代码结构"],
     "逻辑视图按 DDD 划分为用户域、空间域、预约域、考勤域和 AI 域五大领域。开发视图采用 Maven 多模块结构，后端每个服务遵循 Controller、Service、Mapper、Entity 分层，前端采用 Vue 3 组合式 API 与 Pinia 状态管理。",
     "代码结构在 IDEA 中展示更清晰"),
    (9, "4+1 视图 — 物理视图", "03:30 – 03:55",
     ["PPT 第 12 页「物理视图」", "展示 Docker Compose 部署拓扑图"],
     "物理视图展示了 Docker Compose 部署拓扑。系统由 11 个容器组成，统一运行在 campus-network 网络中。流量从前端 nginx 进入网关，再由网关路由到各微服务，最终访问 MySQL 和 Redis。",
     "这是部署章节的前奏"),
    (10, "安全架构", "03:55 – 04:30",
     ["PPT 第 13 页「安全架构」", "可展示登录页面和 JWT 流程"],
     "安全方面，系统采用 JWT + Refresh Token 无状态认证，accessToken 有效期 2 小时，refreshToken 7 天。授权采用 RBAC 模型，分为学生、管理员、超级管理员三级。密码使用 BCrypt 加密，接口通过参数化 SQL 和 XSS 过滤保障数据安全。",
     "可展示登录后 localStorage 中的 token"),
    (11, "数据库设计", "04:30 – 05:00",
     ["PPT 第 14 页「数据库设计」", "展示 ER 图或数据库表结构"],
     "数据库设计包含 18 张表，涵盖用户、角色、自习室、座位、预约、考勤、违规、AI 推荐和知识库等。设计遵循第三范式，统一包含审计字段和逻辑删除字段。同时，系统支持 MySQL 8.0 与国产达梦 8 双数据库兼容。",
     "强调国产化适配"),
    (12, "AI 智能推荐", "05:00 – 05:30",
     ["PPT 第 15 页「AI 智能推荐」", "切换到系统「AI 智能推荐」页面，演示点击获取推荐"],
     "AI 智能推荐是本系统的核心创新点之一。我们采用协同过滤与内容推荐混合算法：首先构建用户-房间交互矩阵，计算用户间余弦相似度；然后加权融合协同过滤与内容过滤分数；最后调用智谱 AI 生成自然语言推荐理由。新用户无历史记录时，系统会退回内容推荐和热门房间兜底。",
     "配合屏幕操作演示，不要只讲不动"),
    (13, "RAG 智能客服", "05:30 – 06:00",
     ["PPT 第 16 页「RAG 智能客服」", "切换到系统「AI 客服」页面，演示提问"],
     "RAG 智能客服基于知识库实现检索增强生成。系统对用户问题进行中文 2-gram 分词，从 knowledge_base 表中检索相关文档，将检索结果拼入 prompt 后调用智谱 GLM-4 生成回答。当 API 不可用时，系统会自动降级为本地关键词匹配，确保服务不中断。",
     "演示时输入一个常见问题，如「如何预约自习室」"),
    (14, "AI 辅助开发", "06:00 – 06:15",
     ["PPT 第 17 页「AI 辅助开发实践」"],
     "在项目开发过程中，我们积极使用 Claude、GitHub Copilot、智谱 GLM-4 等 AI 工具辅助编码，累计记录了七条真实的 AI 辅助开发案例，节省约二十小时以上开发时间。",
     "快速带过，不要占用太多时间"),
    (15, "部署架构", "06:15 – 06:40",
     ["PPT 第 18 页「部署架构」", "展示终端执行 docker compose up -d --build"],
     "系统支持 Docker Compose 一键部署。整个栈包含 11 个容器，通过健康检查和 depends_on 保证启动顺序。我们采用预编译 jar 加轻量 JRE17 镜像，构建时间从数十分钟缩短到 1 至 2 分钟。",
     "可加速播放构建过程"),
    (16, "部署验证结果", "06:40 – 07:10",
     ["PPT 第 19 页「部署验证结果」", "展示 docker compose ps、Nacos 服务列表、登录返回 Token"],
     "部署验证结果显示，11 个容器全部正常运行，7 个微服务成功注册到 Nacos。通过网关测试登录接口，返回 200 并正确签发 JWT Token。前端页面访问正常，AI 服务健康检查状态为 UP，全链路已打通。",
     "这是证明系统跑通的关键镜头，多停留几秒"),
    (17, "部署踩坑与解决方案", "07:10 – 07:30",
     ["PPT 第 20 页「部署踩坑与解决方案」"],
     "部署过程中我们也遇到并解决了多个问题，包括 Docker Hub 拉取超时、Nacos 端口被 Windows 保留、服务配置中心启动崩溃、Redis 配置路径变更，以及网关缺少 LoadBalancer 依赖等。这些问题和解决方案已整理进部署验证报告。",
     "体现工程实践能力"),
    (18, "团队分工", "07:30 – 07:45",
     ["PPT 第 21 页「团队分工」"],
     "团队分工方面，骆家武负责整体架构设计、微服务开发、数据库设计、AI 算法与部署；何展霖负责 Vue3 前端开发、前后端联调、文档撰写、答辩 PPT 和演示视频制作。",
     "根据实际情况填写成员姓名"),
    (19, "项目亮点", "07:45 – 08:10",
     ["PPT 第 22 页「项目亮点与创新点」"],
     "本项目有四大亮点：一是微服务架构完整落地，包含服务拆分、Nacos 注册发现、网关动态路由；二是 AI 双引擎赋能，协同过滤推荐与 RAG 智能客服真实接入智谱 GLM-4；三是支持 MySQL 与国产达梦 8 双数据库；四是 Docker Compose 一键部署并通过全链路验证。",
     "这是答辩加分重点，语速放慢"),
    (20, "总结与展望", "08:10 – 08:35",
     ["PPT 第 23 页「总结与展望」"],
     "目前，系统已完成核心预约、考勤、违规流程，AI 推荐与客服功能，以及 Docker Compose 部署验证。未来可进一步引入 RabbitMQ 实现事件异步解耦，接入 SkyWalking、Prometheus、Grafana 实现全链路可观测，并部署到 Kubernetes 集群实现自动扩缩容。",
     "展望部分简洁有力"),
    (21, "片尾", "08:35 – 09:00",
     ["PPT 第 24 页「致谢」", "淡入系统首页或团队合影", "最后黑场显示「谢谢观看」"],
     "以上就是我们团队关于校园自习室预约系统的全部介绍。感谢指导老师的悉心指导，感谢团队成员的共同努力。欢迎老师和同学们批评指正。",
     "结尾庄重，留给评委反应时间"),
]


# ── 通用样式工具 ────────────────────────────────────────────────
def set_base_font(doc):
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(11)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def style_run(run, size=11, bold=False, color=GRAY):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run(text), size=22, bold=True, color=BLUE)
    return p


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run(text), size=13, bold=False, color=BLUE2)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    style_run(p.add_run(text), size=15, bold=True, color=BLUE)
    return p


def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {qn("w:val"): "clear",
                                         qn("w:color"): "auto",
                                         qn("w:fill"): hexcolor})
    tcPr.append(shd)


def fill_cell(cell, text, bold=False, color=GRAY, size=10.5, white=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    style_run(r, size=size, bold=bold, color=(RGBColor(0xFF, 0xFF, 0xFF) if white else color))


# ════════════════════════════════════════════════════════════════
# 文档一：分镜脚本（无旁白）
# ════════════════════════════════════════════════════════════════
def build_storyboard(path):
    doc = Document()
    set_base_font(doc)

    add_title(doc, f"{TITLE}分镜脚本")
    add_subtitle(doc, "（分镜表 · 不含台词，仅供拍摄/录屏执行）")
    doc.add_paragraph()

    for line in SUBINFO:
        p = doc.add_paragraph()
        style_run(p.add_run("• " + line), size=11, color=GRAY)

    # 时间分配表
    add_h2(doc, "一、整体时间分配")
    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    heads = ["段落", "时间", "内容", "建议时长"]
    for i, h in enumerate(heads):
        shade_cell(t.rows[0].cells[i], "1E3A8A")
        fill_cell(t.rows[0].cells[i], h, bold=True, white=True)
    for row in TIME_TABLE:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            fill_cell(cells[i], v)
    p = doc.add_paragraph()
    style_run(p.add_run("总计：约 9 分钟"), size=11, bold=True, color=BLUE2)

    # 分镜表
    add_h2(doc, "二、分镜表")
    bt = doc.add_table(rows=1, cols=4)
    bt.style = "Table Grid"
    bt.alignment = WD_TABLE_ALIGNMENT.CENTER
    bheads = ["镜头", "时间", "画面内容", "拍摄/录制备注"]
    widths = [Cm(2.0), Cm(3.0), Cm(7.5), Cm(5.0)]
    for i, h in enumerate(bheads):
        shade_cell(bt.rows[0].cells[i], "1E3A8A")
        fill_cell(bt.rows[0].cells[i], h, bold=True, white=True)
    for num, title, tm, screens, _narr, note in SHOTS:
        cells = bt.add_row().cells
        # 镜头号 + 标题
        fill_cell(cells[0], f"{num}", bold=True, color=BLUE2)
        c1 = cells[0].add_paragraph()
        style_run(c1.add_run(title), size=9, bold=False, color=GRAY)
        fill_cell(cells[1], tm)
        # 画面要点（多行）
        cells[2].text = ""
        for j, s in enumerate(screens):
            pp = cells[2].paragraphs[0] if j == 0 else cells[2].add_paragraph()
            style_run(pp.add_run("· " + s), size=10)
        fill_cell(cells[3], note)
    # 设置列宽
    for row in bt.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w

    # 录制建议
    add_h2(doc, "三、录制与后期建议")
    blocks = [
        ("画面录制", [
            "使用 OBS Studio 或 Windows 自带 Xbox Game Bar 录制屏幕",
            "建议录制 PPT 全屏播放 + 系统实际操作画面",
            "关键操作（如 AI 推荐、登录、docker compose ps）可适当放大或加鼠标高亮",
        ]),
        ("后期剪辑", [
            "添加转场动画，但不宜过多",
            "关键页面停留 3–5 秒",
            "操作演示部分保留完整操作流程，不要剪辑过快",
            "背景音乐可选，音量要小，不要盖过旁白",
        ]),
        ("输出规格", [
            "格式：MP4 ／ 分辨率：1920×1080 ／ 帧率：30fps ／ 码率：5–8 Mbps",
            "文件名：校园自习室预约系统_演示视频_v1.mp4",
        ]),
    ]
    for h, items in blocks:
        p = doc.add_paragraph()
        style_run(p.add_run("▪ " + h), size=12, bold=True, color=BLUE2)
        for it in items:
            pp = doc.add_paragraph()
            pp.paragraph_format.left_indent = Cm(0.8)
            style_run(pp.add_run("- " + it), size=10.5)

    # 检查清单
    add_h2(doc, "四、检查清单")
    checks = [
        "视频时长控制在 8–10 分钟", "画面清晰，无黑边或拉伸", "旁白清晰，无杂音",
        "每个技术点都有对应画面支撑",
        "AI 推荐、AI 客服、登录、部署验证等关键场景有实际操作演示",
        "片头片尾完整", "字幕可选，建议添加关键页面字幕",
    ]
    for c in checks:
        p = doc.add_paragraph()
        style_run(p.add_run("☐ " + c), size=10.5)

    doc.save(path)
    return len(SHOTS)


# ════════════════════════════════════════════════════════════════
# 文档二：台词稿（仅旁白）
# ════════════════════════════════════════════════════════════════
def build_script(path):
    doc = Document()
    set_base_font(doc)

    add_title(doc, f"{TITLE}台词稿")
    add_subtitle(doc, "（配音旁白逐字稿 · 按镜头顺序）")
    doc.add_paragraph()

    info = doc.add_paragraph()
    style_run(info.add_run("配音提示："), size=11, bold=True, color=BLUE2)
    style_run(info.add_run("语速控制在每分钟 220–250 字；每段旁白可单独录制，后期拼接；"
                           "用麦克风或耳机录音，避免环境噪音。"), size=11, color=GRAY)
    doc.add_paragraph()

    total_chars = 0
    for num, title, tm, _screens, narr, _note in SHOTS:
        # 小标题：镜头号 + 标题 + 时间
        p = doc.add_paragraph()
        style_run(p.add_run(f"镜头 {num}　{title}"), size=13, bold=True, color=BLUE)
        style_run(p.add_run(f"　（{tm}）"), size=11, bold=False, color=BLUE2)
        # 旁白正文
        body = doc.add_paragraph()
        body.paragraph_format.left_indent = Cm(0.5)
        body.paragraph_format.space_after = Pt(10)
        r = body.add_run(narr)
        style_run(r, size=12, color=GRAY)
        r.font.size = Pt(12)
        total_chars += len(narr)

    # 统计
    doc.add_paragraph()
    stat = doc.add_paragraph()
    est_min = total_chars / 235
    style_run(stat.add_run(
        f"全文约 {total_chars} 字，按每分钟约 235 字配音，预计时长约 {est_min:.1f} 分钟。"),
        size=10.5, bold=True, color=BLUE2)

    doc.save(path)
    return total_chars


if __name__ == "__main__":
    n = build_storyboard("演示视频分镜脚本.docx")
    chars = build_script("演示视频台词稿.docx")
    print(f"OK 分镜脚本.docx -> {n} 个镜头")
    print(f"OK 台词稿.docx   -> 共 {chars} 字旁白")
