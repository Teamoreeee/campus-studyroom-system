#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""根据《2025-2026学年第2学期《软件架构技术》期末综合设计.pdf》评分标准，
对 campus-studyroom 项目进行逐项打分并生成评分报告 Word 文档。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

FONT = "微软雅黑"
BLUE = RGBColor(0x1E, 0x3A, 0x8A)
BLUE2 = RGBColor(0x25, 0x63, 0xEB)
GREEN = RGBColor(0x16, 0xA3, 0x4A)
ORANGE = RGBColor(0xEA, 0x58, 0x0C)
RED = RGBColor(0xDC, 0x26, 0x26)
GRAY = RGBColor(0x1F, 0x29, 0x37)

SCORES = [
    ("需求分析", 10, 8,
     "业务痛点明确（占座、找座难、人工考勤），AI 场景合理（智能推荐+RAG 客服），"
     "需求规格说明书 527 行，覆盖角色权限与核心用例。",
     "可补充量化非功能需求表（响应时间≤250ms、并发≥800、TPS≥400）的实测或设计值。"),
    ("架构设计与多视图建模", 40, 33,
     "4+1 视图完整（逻辑/开发/进程/物理/场景），DDD 五领域，微服务 7 服务拆分，"
     "Spring Cloud Alibaba + Nacos + Gateway，工厂/策略/观察者三种设计模式。",
     "场景视图仅 1 条主线，建议按 rubric 补齐 4 个关键场景（高并发、大数据查询、"
     "服务熔断、AI 推理）的架构应对方案；物理视图可补充节点配置与安全区域。"),
    ("数据库设计与国产化适配", 10, 8,
     "ER 图与 18 张表设计，包含审计字段、逻辑删除、索引优化；"
     "test/sql/dameng/ 下提供达梦 8 兼容脚本。",
     "建议补充国产化适配实测截图；冷热数据分离、读写分离方案可再细化。"),
    ("前后端编码实现", 15, 13,
     "7 个后端微服务 + Vue3 前端全部实现，JWT/RBAC 安全落地，"
     "AI 推荐与 RAG 客服已接入智谱 GLM-4；AI 辅助开发记录 7 条（超过 5 条要求）。"
     "Swagger/OpenAPI 已通过网关聚合全部微服务（79 paths / 12 tags）。",
     "核心功能演示视频目前只有脚本，尚未生成 MP4。"),
    ("云原生部署与测试", 12, 9,
     "Docker Compose 一键部署已跑通，11 个容器 Up，7 个服务注册 Nacos，"
     "全链路登录验证通过；Kubernetes（kind）部署验证已补充截图与终端输出。",
     "JMeter 性能测试目录为空；未做 OWASP ZAP 安全测试；"
     "RabbitMQ 未在 docker-compose 中部署；SkyWalking/Prometheus/Grafana "
     "仅存在于 k8s YAML，未在 Docker Compose 中实际运行。"),
    ("文档与答辩", 13, 12,
     "期末综合设计报告约 14.3 万字（约 70–100 页），答辩 PPT 已生成，"
     "Docker/K8s 系统截图已补齐，演示视频分镜与台词脚本已拆分并分工。",
     "演示视频尚未实际录制。"),
]

AI_BONUS = [
    ("AI 辅助架构设计", 2, 1,
     "AI 辅助开发记录中部分内容涉及架构与代码生成，但未明确形成「AI 生成架构初稿 + 对比优化」的独立记录。"),
    ("基于 RAG 的智能客服", 3, 3,
     "已实现中文 2-gram 检索 + 智谱 GLM-4 生成回答，具备 API 不可用时本地关键词兜底。"),
    ("协同过滤智能推荐算法", 3, 3,
     "已实现用户-房间协同过滤 + 内容推荐融合，并接入大模型生成推荐理由。"),
    ("AI 自动化测试脚本生成", 1, 0,
     "未找到 AI 生成的自动化测试脚本及执行记录。"),
    ("AI 代码审查与优化", 1, 1,
     "AI 辅助开发记录中包含代码优化与 Bug 排查案例，可整理为一份代码审查报告。"),
]


def set_base_font(doc):
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(11)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def style_run(run, size=11, bold=False, color=GRAY):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run(text), size=22, bold=True, color=BLUE)


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run(text), size=12, bold=False, color=BLUE2)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    size = 16 if level == 1 else 13
    style_run(p.add_run(text), size=size, bold=True, color=BLUE)


def add_para(doc, text, bold=False, color=GRAY, size=11):
    p = doc.add_paragraph()
    style_run(p.add_run(text), size=size, bold=bold, color=color)


def add_score_table(doc):
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    headers = ["考核模块", "满分", "得分", "得分率", "主要依据", "扣分/提升点"]
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                style_run(r, size=10, bold=True, color=BLUE)

    base_total = 0
    for name, full, got, reason, gap in SCORES:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = str(full)
        row[2].text = str(got)
        row[3].text = f"{got/full*100:.0f}%"
        row[4].text = reason
        row[5].text = gap
        base_total += got
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    style_run(r, size=9, color=GRAY)

    # total row
    row = table.add_row().cells
    row[0].text = "基础模块合计"
    row[1].text = "90"
    row[2].text = str(base_total)
    row[3].text = f"{base_total/90*100:.1f}%"
    row[4].text = ""
    row[5].text = ""
    for cell in row:
        for p in cell.paragraphs:
            for r in p.runs:
                style_run(r, size=10, bold=True, color=BLUE)

    doc.add_paragraph()
    return base_total


def add_ai_table(doc):
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    headers = ["AI 融合专项加分", "满分", "得分", "主要依据", "扣分/提升点"]
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                style_run(r, size=10, bold=True, color=BLUE)

    ai_total = 0
    for name, full, got, reason in AI_BONUS:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = f"+{full}"
        row[2].text = f"+{got}"
        row[3].text = reason
        row[4].text = "" if got == full else "未完全满足验收标准，建议补齐。"
        ai_total += got
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    style_run(r, size=9, color=GRAY)

    row = table.add_row().cells
    row[0].text = "AI 加分合计"
    row[1].text = "+10"
    row[2].text = f"+{ai_total}"
    row[3].text = ""
    row[4].text = ""
    for cell in row:
        for p in cell.paragraphs:
            for r in p.runs:
                style_run(r, size=10, bold=True, color=BLUE)

    doc.add_paragraph()
    return ai_total


def main():
    doc = Document()
    set_base_font(doc)
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    add_title(doc, "校园自习室预约系统")
    add_subtitle(doc, "期末综合设计评分报告（基于课程评分标准）")

    add_heading(doc, "一、评分说明", level=1)
    add_para(doc,
        "本评分依据《2025-2026学年第2学期〈软件架构技术〉期末综合设计.pdf》的考核模块、"
        "交付物要求与验收标准，对 campus-studyroom 项目当前交付物进行逐项评估。"
        "评分时间为 2026 年 6 月 25 日，评估基于 docs/、backend/、frontend/、test/、k8s/、"
        "docker-compose.yml 等现有文件。",
        size=11)

    add_heading(doc, "二、基础模块评分", level=1)
    base_total = add_score_table(doc)

    add_heading(doc, "三、AI 融合专项加分", level=1)
    ai_total = add_ai_table(doc)

    final = min(base_total + ai_total, 100)
    add_heading(doc, "四、总分", level=1)
    add_para(doc, f"基础模块得分：{base_total} / 90", size=12)
    add_para(doc, f"AI 融合加分：+{ai_total} / +10", size=12)
    add_para(doc, f"最终得分：{final} / 100", bold=True, color=GREEN, size=16)
    add_para(doc,
        "评级：良好（80–89 分段）。项目架构完整、功能实现度高、AI 场景落地扎实，"
        "但云原生部署与测试环节存在明显短板，补齐后有望冲击 90+。",
        size=11)

    add_heading(doc, "五、关键失分点与冲刺建议", level=1)
    suggestions = [
        "JMeter 性能测试：test/jmeter/ 目录为空。建议尽快完成 800 并发、TPS≥400、"
        "平均响应时间≤250ms 的压测，并生成 HTML 报告。",
        "安全测试：使用 OWASP ZAP 扫描一次，确认无高危漏洞，保留扫描报告。",
        "K8s 部署验证：README-k8s.md 中所有验证项均标注为待补充。建议在 Docker Desktop "
        "K8s 上实际跑通，补充 kubectl get pods / services 截图。",
        "运维监控落地：Prometheus/Grafana/SkyWalking 目前只在 k8s YAML 中，未在 "
        "Docker Compose 中运行。答辩视频中需展示监控面板，否则可观测性架构分不足。",
        "RabbitMQ 异步解耦：技术栈要求 Redis + RabbitMQ，但 docker-compose 未部署 "
        "RabbitMQ，代码中可能也未实际使用。",
        "统一接口文档：网关 Swagger 仅路由到 campus-auth，建议通过聚合或各服务独立 "
        "/swagger-ui/index.html 截图补充。",
        "截图与视频：docs/screenshots/ 为空，演示视频尚未录制。这是「文档与答辩」模块的"
        "主要扣分点。",
        "AI 辅助开发记录：已满足 5 条以上，但「AI 自动化测试脚本生成」未做，可补 1 条。",
    ]
    for i, s in enumerate(suggestions, 1):
        p = doc.add_paragraph(style="List Number")
        style_run(p.add_run(s), size=10, color=GRAY)

    add_heading(doc, "六、预计可提升空间", level=1)
    add_para(doc,
        "若补齐上述 8 项（尤其是 JMeter、ZAP、K8s 验证、监控面板、演示视频），"
        "预计可提升 8–12 分，总分有望达到 90–95 分区间。",
        size=11)

    out = "项目评分报告.docx"
    doc.save(out)
    print(f"OK {out}")
    print(f"base={base_total} ai={ai_total} final={final}")


if __name__ == "__main__":
    main()
