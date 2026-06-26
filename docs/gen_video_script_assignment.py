#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""把演示视频台词按团队成员分工标注，生成 演示视频台词稿_分工版.docx。

分工原则：
- 何展霖（前端/文档/PPT/视频）：开场、背景、目标、前端架构、AI 辅助开发、
  团队分工、项目亮点、总结展望、片尾。
- 骆家武（架构/后端/AI/数据）：技术选型、架构论证、微服务拆分、
  4+1 视图（除 shot 8 中前端部分）、数据库、安全、AI 推荐、RAG 客服、
  部署、部署验证、踩坑。
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

FONT = "微软雅黑"
BLUE = RGBColor(0x1E, 0x3A, 0x8A)
BLUE2 = RGBColor(0x25, 0x63, 0xEB)
RED = RGBColor(0xDC, 0x26, 0x26)
GRAY = RGBColor(0x1F, 0x29, 0x37)

# 每个镜头: (序号, 标题, 时间, 旁白, 备注, 配音人)
SHOTS = [
    (1, "片头", "00:00 – 00:30",
     "大家好，我们是 CampusStudio 团队。今天为大家介绍我们的课程设计项目——校园自习室预约系统。本项目基于 Spring Cloud Alibaba 微服务架构，结合 AI 智能推荐与 RAG 智能客服，致力于解决高校自习室资源管理中的痛点问题。",
     "语速适中，营造正式感", "何展霖"),
    (2, "项目背景", "00:30 – 01:00",
     "在高校日常学习场景中，自习室资源管理面临几大痛点：高峰期一座难求，学生常常需要逐层找座；占座不学现象普遍，造成资源浪费；考勤统计依赖人工，效率低且易出错；同时，学生缺乏智能化的座位推荐手段。",
     "照片可网络搜索「大学自习室占座」通用图", "何展霖"),
    (3, "项目目标与需求", "01:00 – 01:30",
     "针对这些痛点，我们设定了五大目标：资源数字化、预约智能化、考勤自动化、管理精细化、架构云原生化。系统覆盖七大业务模块，支持学生、管理员、超级管理员三类角色，并围绕创建预约、签到、AI 推荐等八个核心用例展开设计。",
     "两页切换要快，不要停留太久", "何展霖"),
    (4, "技术选型", "01:30 – 02:00",
     "在技术选型上，后端采用 Java 17 + Spring Boot 3.2.5 + Spring Cloud Alibaba 2023.0.1.2，前端采用 Vue 3 + TypeScript + Element Plus。数据层使用 MySQL 8.0，并兼容国产达梦 8 数据库。缓存用 Redis，服务治理用 Nacos，网关采用 Spring Cloud Gateway，整体支持 Docker 与 Kubernetes 云原生部署。",
     "读技术栈时语速稍快，体现专业性", "骆家武"),
    (5, "架构选型论证", "02:00 – 02:20",
     "我们对比了微服务与单体架构，以及 Spring Cloud Alibaba 与 Netflix 方案。考虑到系统涉及七个业务领域，且需要体现微服务架构能力，最终选择微服务 + Spring Cloud Alibaba，以满足功能实现和国产化信创的双重要求。",
     "强调「信创」这个加分点", "骆家武"),
    (6, "微服务拆分", "02:20 – 02:45",
     "系统拆分为七个微服务：认证服务 auth、用户服务 user、自习室服务 room、预约服务 reservation、考勤服务 attendance、AI 服务 ai，以及统一的 API 网关 gateway。每个服务职责清晰，独立开发、独立部署。",
     "可在 IDEA 中展示 backend 目录下各服务模块", "骆家武"),
    (7, "4+1 视图 — 场景视图", "02:45 – 03:05",
     "架构设计采用 4+1 视图方法。以学生预约自习室为主线：学生登录后浏览自习室，选择座位和时间段提交预约，系统进行冲突检测并创建记录，最后在预约时段内签到签退。",
     "场景视图是串联其他视图的线索", "骆家武"),
    (8, "4+1 视图 — 逻辑视图与开发视图", "03:05 – 03:30",
     "逻辑视图按 DDD 划分为用户域、空间域、预约域、考勤域和 AI 域五大领域。开发视图采用 Maven 多模块结构，后端每个服务遵循 Controller、Service、Mapper、Entity 分层，前端采用 Vue 3 组合式 API 与 Pinia 状态管理。",
     "代码结构在 IDEA 中展示更清晰", "骆家武"),
    (9, "4+1 视图 — 物理视图", "03:30 – 03:55",
     "物理视图展示了 Docker Compose 部署拓扑。系统由 11 个容器组成，统一运行在 campus-network 网络中。流量从前端 nginx 进入网关，再由网关路由到各微服务，最终访问 MySQL 和 Redis。",
     "这是部署章节的前奏", "骆家武"),
    (10, "安全架构", "03:55 – 04:30",
     "安全方面，系统采用 JWT + Refresh Token 无状态认证，accessToken 有效期 2 小时，refreshToken 7 天。授权采用 RBAC 模型，分为学生、管理员、超级管理员三级。密码使用 BCrypt 加密，接口通过参数化 SQL 和 XSS 过滤保障数据安全。",
     "可展示登录后 localStorage 中的 token", "骆家武"),
    (11, "数据库设计", "04:30 – 05:00",
     "数据库设计包含 18 张表，涵盖用户、角色、自习室、座位、预约、考勤、违规、AI 推荐和知识库等。设计遵循第三范式，统一包含审计字段和逻辑删除字段。同时，系统支持 MySQL 8.0 与国产达梦 8 双数据库兼容。",
     "强调国产化适配", "骆家武"),
    (12, "AI 智能推荐", "05:00 – 05:30",
     "AI 智能推荐是本系统的核心创新点之一。我们采用协同过滤与内容推荐混合算法：首先构建用户-房间交互矩阵，计算用户间余弦相似度；然后加权融合协同过滤与内容过滤分数；最后调用智谱 AI 生成自然语言推荐理由。新用户无历史记录时，系统会退回内容推荐和热门房间兜底。",
     "配合屏幕操作演示，不要只讲不动", "骆家武"),
    (13, "RAG 智能客服", "05:30 – 06:00",
     "RAG 智能客服基于知识库实现检索增强生成。系统对用户问题进行中文 2-gram 分词，从 knowledge_base 表中检索相关文档，将检索结果拼入 prompt 后调用智谱 GLM-4 生成回答。当 API 不可用时，系统会自动降级为本地关键词匹配，确保服务不中断。",
     "演示时输入一个常见问题，如「如何预约自习室」", "骆家武"),
    (14, "AI 辅助开发", "06:00 – 06:15",
     "在项目开发过程中，我们积极使用 Claude、GitHub Copilot、智谱 GLM-4 等 AI 工具辅助编码，累计记录了七条真实的 AI 辅助开发案例，节省约二十小时以上开发时间。",
     "快速带过，不要占用太多时间", "何展霖"),
    (15, "部署架构", "06:15 – 06:40",
     "系统支持 Docker Compose 一键部署。整个栈包含 11 个容器，通过健康检查和 depends_on 保证启动顺序。我们采用预编译 jar 加轻量 JRE17 镜像，构建时间从数十分钟缩短到 1 至 2 分钟。",
     "可加速播放构建过程", "骆家武"),
    (16, "部署验证结果", "06:40 – 07:10",
     "部署验证结果显示，11 个容器全部正常运行，7 个微服务成功注册到 Nacos。通过网关测试登录接口，返回 200 并正确签发 JWT Token。前端页面访问正常，AI 服务健康检查状态为 UP，全链路已打通。",
     "这是证明系统跑通的关键镜头，多停留几秒", "骆家武"),
    (17, "部署踩坑与解决方案", "07:10 – 07:30",
     "部署过程中我们也遇到并解决了多个问题，包括 Docker Hub 拉取超时、Nacos 端口被 Windows 保留、服务配置中心启动崩溃、Redis 配置路径变更，以及网关缺少 LoadBalancer 依赖等。这些问题和解决方案已整理进部署验证报告。",
     "体现工程实践能力", "骆家武"),
    (18, "团队分工", "07:30 – 07:45",
     "团队分工方面，骆家武负责整体架构设计、微服务开发、数据库设计、AI 算法与部署；何展霖负责 Vue3 前端开发、前后端联调、文档撰写、答辩 PPT 和演示视频制作。",
     "根据实际情况填写成员姓名", "何展霖"),
    (19, "项目亮点", "07:45 – 08:10",
     "本项目有四大亮点：一是微服务架构完整落地，包含服务拆分、Nacos 注册发现、网关动态路由；二是 AI 双引擎赋能，协同过滤推荐与 RAG 智能客服真实接入智谱 GLM-4；三是支持 MySQL 与国产达梦 8 双数据库；四是 Docker Compose 一键部署并通过全链路验证。",
     "这是答辩加分重点，语速放慢", "何展霖"),
    (20, "总结与展望", "08:10 – 08:35",
     "目前，系统已完成核心预约、考勤、违规流程，AI 推荐与客服功能，以及 Docker Compose 部署验证。未来可进一步引入 RabbitMQ 实现事件异步解耦，接入 SkyWalking、Prometheus、Grafana 实现全链路可观测，并部署到 Kubernetes 集群实现自动扩缩容。",
     "展望部分简洁有力", "何展霖"),
    (21, "片尾", "08:35 – 09:00",
     "以上就是我们团队关于校园自习室预约系统的全部介绍。感谢指导老师的悉心指导，感谢团队成员的共同努力。欢迎老师和同学们批评指正。",
     "结尾庄重，留给评委反应时间", "何展霖"),
]


def set_base_font(doc):
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(11)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def style_run(run, size=11, bold=False, color=GRAY):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run(text), size=22, bold=True, color=BLUE)
    return p


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run(text), size=12, bold=False, color=BLUE2)
    return p


def add_speaker_summary(doc):
    doc.add_paragraph()
    p = doc.add_paragraph()
    style_run(p.add_run("一、配音分工说明"), size=14, bold=True, color=BLUE)

    info = (
        "何展霖（前端负责人 / 文档与答辩负责人）：负责开场、项目背景、目标需求、"
        "AI 辅助开发说明、团队分工、项目亮点、总结展望与片尾。\n"
        "骆家武（架构师 / 后端与 AI 负责人）：负责技术选型、架构论证、微服务拆分、"
        "4+1 视图、数据库设计、安全架构、AI 智能推荐、RAG 智能客服、部署架构、"
        "部署验证与踩坑总结。"
    )
    p = doc.add_paragraph()
    style_run(p.add_run(info), size=11, color=GRAY)

    tips = (
        "录制建议：每人按自己负责模块单独录音，后期按时间轴拼接；"
        "两人音色/语速尽量统一，关键镜头（亮点、部署验证）适当放慢。"
    )
    p = doc.add_paragraph()
    style_run(p.add_run(tips), size=10, color=BLUE2)

    doc.add_paragraph()
    p = doc.add_paragraph()
    style_run(p.add_run("二、分镜头台词稿"), size=14, bold=True, color=BLUE)


def add_shot(doc, num, title, time, narrator, note, speaker):
    speaker_color = BLUE2 if speaker == "骆家武" else RED
    p = doc.add_paragraph()
    style_run(p.add_run(f"镜头 {num}：{title}"), size=13, bold=True, color=BLUE)

    p = doc.add_paragraph()
    style_run(p.add_run(f"时间：{time}    配音："), size=10, color=GRAY)
    style_run(p.add_run(speaker), size=10, bold=True, color=speaker_color)

    p = doc.add_paragraph()
    style_run(p.add_run("旁白："), size=11, bold=True, color=GRAY)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(6)
    style_run(p.add_run(narrator), size=11, color=GRAY)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    style_run(p.add_run(f"备注：{note}"), size=9, color=BLUE2)
    doc.add_paragraph()


def main():
    doc = Document()
    set_base_font(doc)
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    add_title(doc, "校园自习室预约系统")
    add_subtitle(doc, "演示视频台词稿（含配音分工）")
    add_speaker_summary(doc)

    stats = {"骆家武": 0, "何展霖": 0}
    for shot in SHOTS:
        num, title, time, narrator, note, speaker = shot
        add_shot(doc, num, title, time, narrator, note, speaker)
        stats[speaker] += len(narrator)

    doc.add_paragraph()
    p = doc.add_paragraph()
    style_run(p.add_run("三、字数统计"), size=14, bold=True, color=BLUE)

    total = stats["骆家武"] + stats["何展霖"]
    p = doc.add_paragraph()
    style_run(p.add_run(
        f"骆家武：约 {stats['骆家武']} 字（{stats['骆家武'] / total * 100:.0f}%）\n"
        f"何展霖：约 {stats['何展霖']} 字（{stats['何展霖'] / total * 100:.0f}%）\n"
        f"合计：约 {total} 字，按 235 字/分钟约 {total / 235:.1f} 分钟。"
    ), size=11, color=GRAY)

    out = "演示视频台词稿_分工版.docx"
    doc.save(out)
    print(f"OK {out}")
    print(f"  骆家武 {stats['骆家武']} 字 / 何展霖 {stats['何展霖']} 字")


if __name__ == "__main__":
    main()
